import streamlit as st
import pdfplumber
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from google import genai
from supabase import create_client

# ---------------------------------------------------------------------------
# 1. DIRECT DATABASE CONFIGURATION & INITIALIZATION (STABLE DIRECT OVERRIDE)
# ---------------------------------------------------------------------------
st.set_page_config(page_title="SaaS Insurance Portal", layout="wide")

target_url = "https://supabase.co"
target_key = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImtkaWhmbW5zcG55bHlsb2Jycm1rIiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODI3NDExNjksImV4cCI6MjA5ODMxNzE2OX0.kMY9NwQhEI3VAsmLPR3y5XDOZ6FBcsW-0Y-2SOPudjs"

supabase = create_client(target_url, target_key)
supabase.auth._client.base_url = f"{target_url}/auth/v1/"

if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'user_email' not in st.session_state:
    st.session_state['user_email'] = None

if not st.session_state['logged_in']:
    st.title("🛡️ BrokerFlow AI - Insurance Portal")
    st.subheader("Enterprise Proposal Suite")
    
    auth_mode = st.radio("Choose Action", ["Login to Account", "Create Free Agent Account"])
    email = st.text_input("Agency Email Address")
    password = st.text_input("Password", type="password")
    
    if auth_mode == "Create Free Agent Account":
        if st.button("Register as Agent"):
            if len(password) < 6:
                st.error("❌ Password must be at least 6 characters long!")
            else:
                try:
                    res = supabase.auth.sign_up({"email": email, "password": password})
                    st.success("✅ Account created successfully! Please select 'Login to Account' above to log in.")
                except Exception as e:
                    st.error(f"Registration failed: {str(e)}")
                
    elif auth_mode == "Login to Account":
        if st.button("Secure Login"):
            try:
                res = supabase.auth.sign_in_with_password({"email": email, "password": password})
                if res.user:
                    st.session_state['logged_in'] = True
                    st.session_state['user_email'] = res.user.email
                    st.rerun()
            except Exception as e:
                st.error("❌ Invalid agency credentials. Please check your spelling or register.")
    st.stop()
# ---------------------------------------------------------------------------
# 3. HELPER FUNCTIONS FOR WORD DOCUMENT DESIGN
# ---------------------------------------------------------------------------
def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def style_text_run(run, font_name="Calibri", size_pt=10, bold=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb

# ---------------------------------------------------------------------------
# 4. CORE PROCESSING SUITE
# ---------------------------------------------------------------------------
st.sidebar.markdown(f"👤 **Agent Workspace:**\n{st.session_state['user_email']}")
if st.sidebar.button("Logout"):
    st.session_state['logged_in'] = False
    st.session_state['user_email'] = None
    st.rerun()

st.title("📋 AI-Powered Insurance Proposal Generator")
st.caption("Commercial Grade Production Suite v1.0")

# --- SECURE BACKEND API KEY HARDCODING ---
api_key = "AIzaSyCugsZzoIvXYVSZxUDtEPLDLeKoFJW_joA"

client = genai.Client(api_key=api_key)

def extract_text_from_pdf(file_file):
    text = ""
    with pdfplumber.open(file_file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_file):
    doc = docx.Document(file_file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_txt(file_file):
    return file_file.read().decode("utf-8")

st.header("Step 1: Input Raw Data & Documents")
uploaded_files = st.file_uploader(
    "Upload Quote PDFs, Emails (.txt), Word Docs, or Raw Notes", 
    type=["pdf", "docx", "txt"], 
    accept_multiple_files=True
)

additional_notes = st.text_area("Add any additional email text, client requirements, or manual notes here:")
extracted_corpus = ""

if uploaded_files or additional_notes:
    st.info("Documents received! Extracting data fields safely...")
    
    for uploaded_file in uploaded_files:
        if uploaded_file.name.endswith('.pdf'):
            extracted_corpus += f"\n--- Start of File: {uploaded_file.name} ---\n" + extract_text_from_pdf(uploaded_file)
        elif uploaded_file.name.endswith('.docx'):
            extracted_corpus += f"\n--- Start of File: {uploaded_file.name} ---\n" + extract_text_from_docx(uploaded_file)
        elif uploaded_file.name.endswith('.txt'):
            extracted_corpus += f"\n--- Start of File: {uploaded_file.name} ---\n" + extract_text_from_txt(uploaded_file)
            
    if additional_notes:
        extracted_corpus += f"\n--- Start of Additional Notes ---\n{additional_notes}"

    st.header("Step 2: AI Data Compilation & Analysis")
    if st.button("Generate Proposal Analysis Options"):
        with st.spinner("Analyzing coverage lines, structural variances, limits, and premiums..."):
            
            prompt = f"""
            You are an expert commercial insurance broker and underwriter. Analyze the following extracted raw insurance data, quotes, emails, or notes.
            
            Extract ALL insurance quote structural variations and options.
            For each premium tier group/type found (e.g., Primary ABC D&O, Excess Side A Only, Primary Side A Only), format a text block exactly like this so the word parser can map it:
            
            [START_GRID]
            GRID_TITLE: Primary ABC D&O
            OPTION_1: 1 MM | Limit: $1,000,000 | Deductible: Side A $0, Side B $750k | Premium: $90,000
            OPTION_2: 2 MM | Limit: $2,000,000 | Deductible: Side A $0, Side B $750k | Premium: $150,000
            OPTION_3: 3 MM | Limit: $3,000,000 | Deductible: Side A $0, Side B $750k | Premium: $180,000
            [END_GRID]
            
            Provide a clear side-by-side textual analysis summary below the grids explaining the best option for cost vs risk coverage.
            
            Raw Extracted Data:
            {extracted_corpus}
            """
            
            try:
                response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                st.session_state['analysis_result'] = response.text
                st.markdown(response.text)
            except Exception as e:
                st.error(f"Processing Error: {e}")

    if 'analysis_result' in st.session_state:
        st.header("Step 3: Export Proposal")
        
        def create_word_doc(ai_text):
            doc = docx.Document()
            
            sections = doc.sections
            for section in sections:
                new_width, new_height = section.page_height, section.page_width
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = new_width
                section.page_height = new_height
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("COMMERCIAL INSURANCE COVERAGE PROPOSAL OPTIONS")
            style_text_run(run, font_name="Arial", size_pt=14, bold=True, color_rgb=RGBColor(0, 51, 102))
            
            doc.add_heading("Expiring Coverage Summary", level=2)
            exp_table = doc.add_table(rows=2, cols=10)
            exp_table.style = 'Table Grid'
            
            headers = ["Coverage Program", "Opt 1 Limit", "Deductible", "Premium", "Opt 2 Limit", "Deductible", "Premium", "Opt 3 Limit", "Deductible", "Premium"]
            hdr_cells = exp_table.rows[0].cells
            for idx, text in enumerate(headers):
                hdr_cells[idx].text = text
                set_cell_background(hdr_cells[idx], "003366")
                set_cell_margins(hdr_cells[idx])
                for para in hdr_cells[idx].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in para.runs:
                        style_text_run(r, font_name="Arial", size_pt=9, bold=True, color_rgb=RGBColor(255, 255, 255))
            
            doc.add_paragraph()
            doc.add_heading("Proposed Coverage Marketing Options", level=2)
            
            lines = ai_text.split('\n')
            in_grid = False
            grid_title = "Proposed Options"
            options_extracted = []
            
            for line in lines:
                if "[START_GRID]" in line:
                    in_grid = True
                    options_extracted = []
                    continue
                if "[END_GRID]" in line:
                    in_grid = False
                    
                    table = doc.add_table(rows=3, cols=10)
                    table.style = 'Table Grid'
                    
                    r0 = table.rows[0].cells
                    r0[0].text = grid_title
                    set_cell_background(r0[0], "003366")
                    
                    col_offset = 1
                    for idx, opt in enumerate(options_extracted):
                        if idx >= 3: break
                        start_cell = r0[col_offset]
                        end_cell = r0[col_offset + 2]
                        merged_hdr = start_cell.merge(end_cell)
                        merged_hdr.text = f"{grid_title} - {opt.get('name', '')}"
                        set_cell_background(merged_hdr, "0066CC")
                        col_offset += 3
                        
                    r1 = table.rows[1].cells
                    r1[0].text = "Aggregate Limit Structure"
                    set_cell_background(r1[0], "F2F2F2")
                    
                    sub_hdrs = ["Limit", "Deductible", "Premium"] * 3
                    for i, sh in enumerate(sub_hdrs):
                        r1[i+1].text = sh
                        set_cell_background(r1[i+1], "E6F2FF")
                        
                    r2 = table.rows[2].cells
                    r2[0].text = "Maximum Aggregate Limit of Liability\nTotal Premium + Fees & Taxes"
                    
                    data_offset = 1
                    for opt in options_extracted:
                        r2[data_offset].text = opt.get('limit', '')
                        r2[data_offset + 1].text = opt.get('deductible', '')
                        r2[data_offset + 2].text = opt.get('premium', '')
                        
                        p_cell = r2[data_offset + 2]
                        for paragraph in p_cell.paragraphs:
                            for run in paragraph.runs:
                                style_text_run(run, font_name="Arial", size_pt=9, bold=True, color_rgb=RGBColor(204, 0, 0))
                        data_offset += 3
                        
                    for row in table.rows:
                        for cell in row.cells:
                            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
                    doc.add_paragraph()
                    continue
                    
                if in_grid:
                    if "GRID_TITLE:" in line:
                        grid_title = line.replace("GRID_TITLE:", "").strip()
                    elif "OPTION_" in line:
                        parts = line.split('|')
                        opt_name = parts[0].split(':')[1].strip() if len(parts) > 0 and ':' in parts[0] else "Option"
                        limit_val = parts[1].replace("Limit:", "").strip() if len(parts) > 1 else ""
                        ded_val = parts[2].replace("Deductible:", "").strip() if len(parts) > 2 else ""
                        prem_val = parts[3].replace("Premium:", "").strip() if len(parts) > 3 else ""
                        
                        options_extracted.append({
                            "name": opt_name, 
                            "limit": limit_val, 
                            "deductible": ded_val, 
                            "premium": prem_val
                        })
                else:
                    if line.strip() and not line.startswith('['):
                        doc.add_paragraph(line)
                        
            output_path = "Premium_Structured_Proposal.docx"
            doc.save(output_path)
            return output_path

        if st.button("Build Word Document"):
            doc_file_path = create_word_doc(st.session_state['analysis_result'])
            with open(doc_file_path, "rb") as f:
                st.download_button(
                    label="📥 Download Word Proposal",
                    data=f,
                    file_name="Premium_Insurance_Proposal.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
